"""Extract raw text from a .docx file using python-docx.

Lazy import — python-docx pulls in lxml which is heavy at boot.
"""
from __future__ import annotations
from io import BytesIO


def extract_text(data: bytes) -> str:
    from docx import Document  # lazy — heavy import
    doc = Document(BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)
